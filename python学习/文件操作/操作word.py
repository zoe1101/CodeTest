from docx import Document, shared


def merge_without_format(docx_files: list):
    '''
    只获取内容进行合并
    '''
    doc = Document()
    # 遍历每个文件
    for docx_file in sorted(docx_files):
        another_doc = Document(docx_file)
        # 获取每个文件的所有“段落”
        paras = another_doc.paragraphs
        # 获取所有段落的文字内容
        # paras_content = [para.text for para in paras]
        for para in paras:
            # 为新的word文件创建一个新段落
            newpar = doc.add_paragraph('')
            # 将提取的内容写入新的文本段落中
            newpar.add_run(para.text)
        # 所有文件合并完成后在指定路径进行保存
    # 在合并后的word中添加图片
    doc.add_picture('test.png', width=shared.Inches(1)) #按英寸设置宽度,添加图片
    doc.save('merge_word.docx')


if __name__ == '__main__':
    merge_without_format(['../../data/word1.docx', '../../data/word2.docx'])
